from dataclasses import replace
from datetime import timedelta
from re import U
import time
import requests 
from selenium import webdriver
import parsel
import re

#获取日期
t=time.gmtime()
tim=time.strftime("%Y/%m/%d",t)
print(tim)
#模拟点击
'''options=webdriver.ChromeOptions()
browser=webdriver.Chrome(options=options)
url_reserch="https://news.cctv.com/?spm=C96370.PPDB2vhvSivD.E59hodVIdh2C.3"
browser.get(url_reserch)
browser.find_element_by_class_name("btn_icon").click()
browser.find_element_by_class_name("input_txt2").send_keys("新闻联播")
browser.find_element_by_id("mytxtdafdfasdf").send_keys("新闻联播")
browser.find_element_by_id("search_btn").click()'''

url_list='https://tv.cctv.com/lm/xwlb/'
header={
    "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/101.0.4951.64 Safari/537.36 Edg/101.0.1210.47"
}
response_list=requests.get(url=url_list,headers=header)
response_list.encoding='UTF-8'
list=response_list.text

#content > li:nth-child(1) > div > a :: attr(href)
try:
    real_url=re.findall('<a href="(.*?)" target="_blank"><i class="sql0">完整版</i>《新闻联播》 20220528 19:00</a>',list)[0].split("target")[0]
    #print(real_url)
    if tim in real_url:
        url=str(real_url.replace('"',''))
        print("中了,不用改了")
except:
    url=input("请输入{}的新闻联播链接:".format(tim))
header={
    "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/101.0.4951.64 Safari/537.36 Edg/101.0.1210.47"
}
response=requests.get(url=url,headers=header)
response.encoding='UTF-8'
content=response.text
selector=parsel.Selector(content)
text1=selector.css('#page_body > div.allcontent > div.video18847 > div.playingCon > div.nrjianjie_shadow > div > ul > li:nth-child(1) > p::text').get()#[0]
print(text1)
flag=eval(input("文章是否紧凑，1为紧凑，0为不紧凑:"))
if( flag == 1 ):
    text2=text1.split("； ")
    #text2=text1[0].replace('\n','').split("；")
    print(text2)
    print(1)
    text='；\n'.join(text2)
else:#这种情况较多，主要按0
    text=text1.replace('\n','')
#title=selector.css('#chbox01 > div.mtab_con > div:nth-child(1) > div > div.text_box_02 > p:nth-child(2) ::text').getall()
title1=selector.css('#page_body > div.allcontent > div.video18847 > div.playingVideo > div.tit ::text').getall()[0][6:15]
title=title1[1:5]+'年'+title1[5:7]+'月'+title1[7:]+'日'

#txt格式写入
'''with open('D:/桌面/'+'每日新闻联播'+'.txt',mode='w',encoding='utf-8') as f:
    f.write(title)
    f.write('\n')
    f.write(text)
    f.write('\n')    
    print("已完成")
'''

import docx
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor

#将文本写入word
file=docx.Document()

#设置正文字体
file.styles['Normal'].font.name = u'宋体'
file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
file.styles['Normal'].font.size = Pt(12)
file.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
#设置标题字体
Head = file.add_heading("",level=3)# 这里不填标题内容，设置3级标题
run  = Head.add_run(title)
run.font.name=u'times new roman'#u'Cambria'
run.font.color.rgb = RGBColor(0,0,0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'times new roman')

#file.add_heading(title, 3)  上述为复杂化
file.add_paragraph(text) 
file.save("D:\桌面\新闻.docx") 
print("{}新闻联播已成功爬取。".format(title))
